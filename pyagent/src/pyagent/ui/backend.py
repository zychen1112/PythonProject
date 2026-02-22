"""
Backend module for PyAgent UI.
Encapsulates all backend capabilities: Hooks, RAG, Memory, State.
"""

from __future__ import annotations

import asyncio
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

from pyagent.core.agent import Agent, AgentConfig
from pyagent.core.tools import Tool
from pyagent.hooks import HookRegistry, LoggingHook, TimingHook, ErrorHandlingHook
from pyagent.memory import (
    ConversationMemory,
    SemanticMemory,
    EpisodicMemory,
    MemoryManager,
    Episode,
)
from pyagent.rag import (
    Document,
    MemoryVectorStore,
    FakeEmbedding,
    RAGPipeline,
    VectorRetriever,
)
from pyagent.state import CheckpointManager, MemoryCheckpointBackend


class PyAgentBackend:
    """
    Backend manager that encapsulates all agent capabilities.

    This class manages:
    - Agent creation with Hooks, Memory, RAG
    - Document indexing (RAG)
    - Memory management
    - State persistence
    """

    def __init__(self):
        self.agent: Optional[Agent] = None
        self.provider: Any = None
        self.config: Optional[AgentConfig] = None

        # RAG components
        self.rag_pipeline: Optional[RAGPipeline] = None
        self.indexed_documents: list[dict[str, Any]] = []

        # Memory components
        self.memory_manager: Optional[MemoryManager] = None

        # State components
        self.checkpoint_manager: CheckpointManager = CheckpointManager(
            MemoryCheckpointBackend()
        )
        self.current_thread_id: str = f"thread_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # Hooks registry
        self.hooks_registry: Optional[HookRegistry] = None

        # Current settings
        self.current_settings: dict[str, Any] = {}

    def initialize(
        self,
        provider_name: str,
        provider_class: type,
        api_key: str,
        model: str,
        temperature: float,
        max_tokens: int,
        system_prompt: str,
    ) -> None:
        """Initialize the backend with all capabilities."""

        # Create provider
        self.provider = provider_class(api_key=api_key)

        # Create config
        self.config = AgentConfig(
            name="PyAgentPro",
            model=model,
            temperature=temperature,
            max_tokens=max_tokens,
            system_prompt=system_prompt if system_prompt.strip() else None,
        )

        # Initialize RAG pipeline
        vectorstore = MemoryVectorStore()
        embedding = FakeEmbedding()  # Use fake embedding for simplicity
        retriever = VectorRetriever(embedding, vectorstore)
        self.rag_pipeline = RAGPipeline(
            embedding=embedding,
            vectorstore=vectorstore,
            retriever=retriever,
        )

        # Initialize Memory Manager
        self.memory_manager = MemoryManager(
            conversation_memory=ConversationMemory(),
            semantic_memory=SemanticMemory(vectorstore, embedding),
            episodic_memory=EpisodicMemory(vectorstore, embedding),
        )

        # Initialize Hooks Registry with built-in hooks
        self.hooks_registry = HookRegistry()
        self.hooks_registry.register(LoggingHook(level="INFO"))
        self.hooks_registry.register(TimingHook(warn_threshold_ms=3000))
        self.hooks_registry.register(ErrorHandlingHook(max_retries=2))

        # Create RAG tool
        rag_tool = self._create_rag_tool()

        # Create demo tools
        demo_tools = self._create_demo_tools()

        # Create Agent
        self.agent = Agent(
            provider=self.provider,
            config=self.config,
            tools=demo_tools + [rag_tool],
            hooks_registry=self.hooks_registry,
        )

        # Store settings
        self.current_settings = {
            "provider_name": provider_name,
            "api_key": api_key,
            "model": model,
            "temperature": temperature,
            "max_tokens": max_tokens,
            "system_prompt": system_prompt,
        }

    def _create_demo_tools(self) -> list[Tool]:
        """Create demo tools for the agent."""

        def get_weather(location: str) -> str:
            return f"{location}今天是晴天，气温25°C，空气质量良好。"

        def calculate(expression: str) -> str:
            try:
                # Safe evaluation for simple math
                import ast
                import operator
                ops = {
                    ast.Add: operator.add,
                    ast.Sub: operator.sub,
                    ast.Mult: operator.mul,
                    ast.Div: operator.truediv,
                    ast.Pow: operator.pow,
                }
                node = ast.parse(expression, mode='eval')
                def _eval(node):
                    if isinstance(node, ast.Expression):
                        return _eval(node.body)
                    elif isinstance(node, ast.Constant):
                        return node.value
                    elif isinstance(node, ast.BinOp):
                        left = _eval(node.left)
                        right = _eval(node.right)
                        return ops[type(node.op)](left, right)
                    else:
                        raise ValueError("Unsupported operation")
                result = _eval(node)
                return f"计算结果: {result}"
            except Exception as e:
                return f"计算错误: {e}"

        def get_current_time() -> str:
            return f"当前时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        return [
            Tool(
                name="get_weather",
                description="获取指定城市的天气信息",
                input_schema={
                    "type": "object",
                    "properties": {
                        "location": {"type": "string", "description": "城市名称，如：北京"}
                    },
                    "required": ["location"]
                },
                handler=get_weather,
            ),
            Tool(
                name="calculate",
                description="计算数学表达式",
                input_schema={
                    "type": "object",
                    "properties": {
                        "expression": {"type": "string", "description": "数学表达式，如：2+3*4"}
                    },
                    "required": ["expression"]
                },
                handler=calculate,
            ),
            Tool(
                name="get_current_time",
                description="获取当前日期和时间",
                input_schema={"type": "object", "properties": {}, "required": []},
                handler=get_current_time,
            ),
        ]

    def _create_rag_tool(self) -> Tool:
        """Create RAG search tool."""

        async def search_knowledge(query: str) -> str:
            if not self.rag_pipeline:
                return "知识库未初始化"
            results = await self.rag_pipeline.retrieve(query, k=3)
            if not results:
                return "未找到相关知识"
            context = "\n\n".join([
                f"[文档 {i+1}]: {r.chunk.content}"
                for i, r in enumerate(results)
            ])
            return f"从知识库检索到以下内容：\n\n{context}"

        return Tool(
            name="search_knowledge",
            description="搜索已上传的知识文档，获取相关信息。当用户询问上传文档的内容时使用此工具。",
            input_schema={
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "搜索关键词或问题"}
                },
                "required": ["query"]
            },
            handler=search_knowledge,
        )

    async def index_document(self, file_path: str, file_name: str) -> str:
        """Index a document into the RAG system."""
        if not self.rag_pipeline:
            return "RAG系统未初始化"

        try:
            # Read file content
            content = self._read_file(file_path)
            if not content:
                return f"无法读取文件: {file_name}"

            # Create document
            doc = Document(
                id=f"doc_{len(self.indexed_documents)}",
                content=content,
                metadata={"source": file_name, "indexed_at": datetime.now().isoformat()}
            )

            # Index document
            await self.rag_pipeline.index([doc])

            # Track indexed document
            self.indexed_documents.append({
                "name": file_name,
                "id": doc.id,
                "size": len(content),
                "chunks": len(content) // 500 + 1,  # Estimate
            })

            return f"文档 '{file_name}' 已成功索引 (约 {len(content)} 字符)"

        except Exception as e:
            return f"索引失败: {str(e)}"

    def _read_file(self, file_path: str) -> str:
        """Read file content based on file type."""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in [".txt", ".md", ".py", ".json", ".yaml", ".yml", ".xml"]:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        elif suffix == ".pdf":
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                return "[PDF解析需要安装pypdf: pip install pypdf]"

        elif suffix in [".docx", ".doc"]:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                return "[Word解析需要安装python-docx: pip install python-docx]"

        else:
            # Try to read as text
            try:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            except:
                return ""

    async def chat(self, message: str) -> str:
        """Process a chat message."""
        if not self.agent:
            return "Agent未初始化，请先配置API Key"

        try:
            response = await self.agent.run(message)

            # Auto-store episode to memory
            if self.memory_manager and self.memory_manager.episodic:
                episode = Episode(
                    id="",
                    summary=f"用户提问: {message[:50]}...",
                    details=message,
                    outcome=response[:200] if response else "",
                )
                await self.memory_manager.episodic.record(episode)

            return response

        except Exception as e:
            return f"错误: {str(e)}"

    async def save_checkpoint(self) -> str:
        """Save current conversation state."""
        if not self.agent:
            return "没有活动的会话"

        try:
            state = {
                "messages": [
                    {"role": msg.role.value, "content": str(msg.content)}
                    for msg in self.agent.context.messages
                ],
                "indexed_documents": self.indexed_documents,
                "settings": self.current_settings,
            }

            checkpoint = await self.checkpoint_manager.save(
                thread_id=self.current_thread_id,
                state=state,
                metadata={"timestamp": datetime.now().isoformat()}
            )

            return f"会话已保存 (ID: {checkpoint.id[:8]}...)"

        except Exception as e:
            return f"保存失败: {str(e)}"

    async def load_checkpoint(self) -> str:
        """Load the latest checkpoint."""
        try:
            checkpoint = await self.checkpoint_manager.get_latest(self.current_thread_id)
            if not checkpoint:
                return "没有找到已保存的会话"

            # Restore state
            state = checkpoint.state
            if self.agent and "messages" in state:
                # Note: Full restoration would require rebuilding messages
                # For now, just return status
                return f"已加载会话 ({len(state.get('messages', []))} 条消息)"

            return "会话已加载"

        except Exception as e:
            return f"加载失败: {str(e)}"

    def clear_memory(self) -> str:
        """Clear all memory and context."""
        if self.agent:
            self.agent.clear_history()

        if self.memory_manager:
            # Clear conversation memory
            if self.memory_manager.conversation:
                self.memory_manager.conversation._entries.clear()

        return "记忆已清除"

    def reset(self) -> str:
        """Reset everything."""
        self.agent = None
        self.indexed_documents = []

        if self.memory_manager:
            self.memory_manager.conversation._entries.clear()
            self.memory_manager.semantic._items.clear()
            self.memory_manager.episodic._episodes.clear()

        # Create new thread ID
        self.current_thread_id = f"thread_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        return "会话已重置"

    def get_status(self) -> dict[str, Any]:
        """Get current status."""
        return {
            "initialized": self.agent is not None,
            "documents": len(self.indexed_documents),
            "document_list": [d["name"] for d in self.indexed_documents],
            "thread_id": self.current_thread_id,
        }


# Global backend instance
_backend: Optional[PyAgentBackend] = None


def get_backend() -> PyAgentBackend:
    """Get or create the global backend instance."""
    global _backend
    if _backend is None:
        _backend = PyAgentBackend()
    return _backend
